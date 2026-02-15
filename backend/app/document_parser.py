"""Document parsing utilities for extracting text from various file formats."""
import io
from typing import Optional

from fastapi import UploadFile


async def extract_text_from_file(file: UploadFile) -> Optional[str]:
    """
    Extract text content from uploaded file.

    Supports: .txt, .pdf, .doc, .docx
    """
    if not file or not file.filename:
        return None

    filename = file.filename.lower()

    # Read content and reset file pointer
    content = await file.read()
    await file.seek(0)  # Reset file pointer for potential re-reads

    if not content:
        print(f"Warning: Empty file content for {filename}")
        return None

    print(f"Processing file: {filename}, size: {len(content)} bytes")

    try:
        if filename.endswith('.txt'):
            result = extract_from_txt(content)
        elif filename.endswith('.pdf'):
            result = extract_from_pdf(content)
        elif filename.endswith('.docx'):
            result = extract_from_docx(content)
        elif filename.endswith('.doc'):
            result = extract_from_doc(content)
        else:
            print(f"Unsupported file format: {filename}")
            return None

        if result:
            print(f"Extracted {len(result)} characters from {filename}")
        return result
    except Exception as e:
        print(f"Error extracting text from {filename}: {e}")
        return None


def extract_from_txt(content: bytes) -> str:
    """Extract text from .txt file."""
    # Try different encodings
    for encoding in ['utf-8', 'utf-16', 'latin-1', 'cp1252']:
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode('utf-8', errors='ignore')


def extract_from_pdf(content: bytes) -> str:
    """Extract text from .pdf file."""
    try:
        from PyPDF2 import PdfReader

        pdf_file = io.BytesIO(content)
        reader = PdfReader(pdf_file)

        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)

        return '\n'.join(text_parts)
    except ImportError:
        return "[PDF parsing requires PyPDF2. Please install it with: pip install PyPDF2]"
    except Exception as e:
        return f"[Error reading PDF: {str(e)}]"


def extract_from_docx(content: bytes) -> str:
    """Extract text from .docx file."""
    try:
        from docx import Document

        docx_file = io.BytesIO(content)
        doc = Document(docx_file)

        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        return '\n'.join(text_parts)
    except ImportError:
        return "[DOCX parsing requires python-docx. Please install it with: pip install python-docx]"
    except Exception as e:
        return f"[Error reading DOCX: {str(e)}]"


def extract_from_doc(content: bytes) -> str:
    """
    Extract text from .doc file.
    Note: .doc format is proprietary and harder to parse without specialized libraries.
    """
    # Try to extract as plain text (works for some .doc files)
    try:
        text = content.decode('utf-8', errors='ignore')
        # Filter out binary garbage
        printable_text = ''.join(char for char in text if char.isprintable() or char in '\n\r\t')
        if len(printable_text) > 50:  # If we got meaningful text
            return printable_text
    except:
        pass

    return "[.doc format has limited support. Please convert to .docx or .txt for better results]"
